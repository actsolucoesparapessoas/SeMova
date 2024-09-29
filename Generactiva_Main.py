import streamlit as st
from streamlit_chat import message
from PIL import Image # Lib para carregar imagem no Streamlit
import matplotlib.pyplot as plt
import pandas as pd
import io
import os
from gtts import gTTS #Lib para Conversão Text2Voice. Em seguida pode usar Gemini para converter voice para texto
import google.generativeai as genai
from openai import OpenAI

import urllib3
from urllib3 import request
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os.path
from wordcloud import WordCloud, STOPWORDS
#NLP Package
from enelvo.normaliser import Normaliser
norm = Normaliser(tokenizer='readable')
GeraGrafico = False #Variável booleana para gerar nuvem de palavras apenas após ter resposta.

import sqlite3

# Cria uma conexão com o banco de dados SQLite3
conn = sqlite3.connect('User.db')
c = conn.cursor()

# Cria a tabela 'User' se ela não existir
c.execute('''CREATE TABLE IF NOT EXISTS User (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    nome TEXT,
    keyGEMINI TEXT,
    keyOPENAI TEXT,
    situacao TEXT
)''')

# Exibe todos os registros da tabela
def mostrar_registros():
    c.execute('''SELECT * FROM User''')
    registros = c.fetchall()
    if registros:
        df = pd.DataFrame(registros, columns=['ID', 'NOME', 'CHAVE_GEMINI', 'CHAVE_OPENAI', 'SITUACAO'])
        #st.dataframe(df)
        CHAVE_GEMINI = df['CHAVE_GEMINI'][0]
        CHAVE_OPENAI = df['CHAVE_OPENAI'][0]
        return CHAVE_GEMINI, CHAVE_OPENAI
    else:
        return '0000000000000000000000000000000000000'
        st.write('Não há registros no banco de dados.')

def api_openai(prompt):   
    OPENAI_KEY = mostrar_registros()[1]
    client = OpenAI(api_key=OPENAI_KEY)
    completion = client.chat.completions.create(
      model="gpt-3.5-turbo",
      messages=[
        {"role": "user", "content": str(prompt)}
      ]
    )
    return completion.choices[0].message.content
    
def api_gemini(QUESTION):
    GEMINI_KEY = mostrar_registros()[0]
    genai.configure(api_key=GEMINI_KEY)

    # Set up the model
    generation_config = {
      "temperature": 0.9,
      "top_p": 1,
      "top_k": 1,
      "max_output_tokens": 2048,
    }
    model = genai.GenerativeModel('gemini-pro')
    response = model.generate_content(QUESTION)
    resp = response.text
    stopwords = set(STOPWORDS)
    stopwords.update(["de", "ao", "o", "nao", "para", "da", "meu", "em", "você", "ter", "um", "ou", "os", "ser", "só"])
    # gerar uma wordcloud
    wordcloud = WordCloud(stopwords=stopwords,
                          background_color="white",
                          width=1280, height=720).generate(resp)
    #resp = resp.replace("*", "")
    #resp = resp.replace("**", "")
    return resp, wordcloud

def Ler_DOCx(file_path):
    DOC = Document(file_path)
    TEXTO = ''
    #TITULO = st.text_input("TÍTULO: ", DOC.paragraphs[0].text)
    Titulo = DOC.paragraphs[0].text
    n = len(DOC.paragraphs)
    for i in range(1, n):
        TEXTO+=DOC.paragraphs[i].text + '\n'
        #st.write(DOC.paragraphs[i].text)
    return Titulo, TEXTO
    
st.set_page_config(layout="wide", page_title="Generactiva v0.1 (By. Massaki)")
image = Image.open('Logov1.png')
st.sidebar.image(image, width=300)
audio_file0 = open('Generactiva.mp3', 'rb')
audio_bytes0 = audio_file0.read()
st.sidebar.audio(audio_bytes0, format='audio/ogg',start_time=0)
Robo = "imgs/stuser.png"
with st.chat_message("user", avatar=Robo):
    st.write("Versão Beta (Em teste)")



Titulo_Principal = '<p style="font-weight: bolder; color:#f55050; font-size: 48px;">GENERACTIVA - V0.1</p>'    
st.markdown(Titulo_Principal, unsafe_allow_html=True)
mystyle0 = '''<style> p{text-align:center;}</style>'''
st.markdown(mystyle0, unsafe_allow_html=True) 
Sub_Titulo = '<p style="font-weight: bolder; color:White; font-size: 18px;">Mais que uma IA Generativa, uma assistente!</p>'
st.markdown(Sub_Titulo, unsafe_allow_html=True)        
st.markdown(mystyle0, unsafe_allow_html=True) 

tab1, tab2, tab3 = st.tabs(["Instruções", "Perguntas&Respostas", "Painel"])
with tab1:
    Sub_Titulo = '<p style="font-weight: bolder; color:#f55050; font-size: 22px;">Dicas de utilização em 4 passos:</p>'
    st.markdown(Sub_Titulo, unsafe_allow_html=True)
    mystyle1 =   '''<style> p{text-align:left;}</style>'''
    st.markdown(mystyle1, unsafe_allow_html=True) 
    st.markdown('''     
                **- Passo 1:**  Faça sua Pesquisa (na aba "Perguntas&Respostas")
                
                **- Passo 2:**  Será exibida a resposta à sua pesquisa (com uma cópia na caixa de texto caso deseje editá-la). 
                
                **- Passo 3:**  Será exibido o Botão ✔️ Salvar resp no arq.DOCx junto ao Histórico de Pesquisa(s) e Resposta(s)
                            
                        OBS: ao clicar no botão ✔️ Salvar resp no arq.DOCx você salvará a resposta obtida e eventualmente editada.
                            
                        Também é possível modificar o título do arq.DOCx a ser salvo clicando na caixa de ↪️ Título do arq.DOCX na bara laterial esquerda.
                
                **- Passo 4:**  Edição do Documento.DOCx salvo anteriormente
                            
                            - Observa-se que para editar o arq.DOCx salvo basta ir na Barra laterial esquerda e clicar no Menu 📝 Editor DOCx
                ''')
    st.divider()
    
    Instrucao = '<p style="font-weight: bolder; color:white; font-size: 26px;">Após ler o passo a passo nesta aba "Instruções", basta clicar na Aba "Perguntas&Respostas".</p>'
    st.markdown(Instrucao, unsafe_allow_html=True)
    st.markdown(mystyle1, unsafe_allow_html=True) 
    
with tab2:     
    Passo1 = '<p style="font-weight: bolder; color:White; font-size: 16px;">Passo 1:</p>'
    st.markdown(Passo1, unsafe_allow_html=True)
    st.markdown(mystyle1, unsafe_allow_html=True) 
    question = norm.normalise(st.text_input("Digite sua pergunta aqui 👇", key="input"))
        
    if 'ai_answer' not in st.session_state:
        st.session_state['ai_answer'] = []

    if 'ai_question' not in st.session_state:
        st.session_state['ai_question'] = []

    if question:
        output = api_gemini(question)[0]
        output2 = api_openai(question)
        NuvemPalavras = api_gemini(question)[1]
        GeraGrafico = True        
        output = output.lstrip("\n")
        output2 = output2.lstrip("\n")
        # Store the outputs
        st.session_state.ai_question.append(question)
        st.session_state.ai_answer.append("**Resposta do Gemini:** \n" + output + " \n " + " \n **Resposta OPENAI:** \n" + output2)
        NomeArq = st.sidebar.text_input("Digite NOME do arquivo.DOCx e tecle ENTER: 👇", 'generactiva.docx')
        TITULO = st.sidebar.text_input("↪️ Título do arq.DOCX: ", "Programado por Massaki Igarashi")
        check_file = os.path.isfile(NomeArq)
        st.sidebar.write(check_file)        
        
        if check_file:
            #TITULO2 = st.sidebar.text_input("TÍTULO: ", Ler_DOCx(NomeArq)[0])
            memo1 = st.sidebar.text_area("Conteúdo: ", Ler_DOCx(NomeArq)[1])
            bio = io.BytesIO()
            #st.sidebar.download_button(label="⬇️ Download do arq.DOCx",
            #                                    data=bio.getvalue(),
            #                                    file_name='https://github.com/actsolucoesparapessoas/Generactiva/blob/master/generactiva.docx',
            #                                    mime="docx")
            with open(NomeArq, "rb") as doc_file:
                 DOC = doc_file.read()
            st.sidebar.download_button(label="⬇️ Download do arq.DOCx",
                                       data=DOC,
                                       file_name=NomeArq,
                                       mime="docx")
        else:
            TITULO2 = NomeArq
 
        
    message_history = st.empty()

    

    if st.session_state['ai_answer']:
        memo = str(st.session_state['ai_answer'][len(st.session_state['ai_answer'])-1])
        Passo2 = '<p style="font-weight: bolder; color:White; font-size: 16px;">Passo 2:</p>'
        st.markdown(Passo2, unsafe_allow_html=True)
        st.markdown(mystyle1, unsafe_allow_html=True)
        st.write(memo)
        memo = memo.replace("*", " ")
        memo2 = st.text_input("Resposta editável:", memo)
        
        Passo3 = '<p style="font-weight: bolder; color:White; font-size: 16px;">Passo 3:</p>'
        st.markdown(Passo3, unsafe_allow_html=True)
        st.markdown(mystyle1, unsafe_allow_html=True)
        #NomeArq = st.sidebar.text_input("Digite nome do arquivo.DOCx e tecle ENTER: 👉", 'generactiva.docx')
        #check_file = os.path.isfile(NomeArq)
        #st.sidebar.write(check_file)
        
        if st.sidebar.button(label = '✔️ Salvar o arquivo .DOCx'):
            if check_file:
                #TITULO2 = st.sidebar.text_input("TÍTULO: ", Ler_DOCx(NomeArq)[0])
                #memo1 = st.sidebar.text_area("Conteúdo: ", Ler_DOCx(NomeArq)[1])
                document = Document()
                #document.add_heading(TITULO2, 0)            
                document.add_heading(TITULO, 0)       
                p = document.add_paragraph(memo1+memo2)
                p.bold = True
                p.italic = True
                #p.add_run('bold').bold = True
                #p.add_run(' and some ')
                #p.add_run('italic.').italic = True
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                document.save(NomeArq)
            else:            
                document = Document()
                document.add_heading(TITULO, 0)              
                p = document.add_paragraph(memo2)
                p.bold = True
                p.italic = True
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                document.save(NomeArq)

 
    if st.session_state['ai_answer']:
        my_expander = st.expander(label='Clique e abra o hitórico de Perguntas & Respostas! 👉')
        with my_expander:
            for i in range(len(st.session_state['ai_answer']) - 1, -1, -1):
                # This function displays Gemini response
                #Exemplos de avatar_styles: https://discuss.streamlit.io/t/what-are-the-other-avatar-style-icon-options/61528/4
                message("PERGUNTA: " + st.session_state['ai_question'][i],                        
                    #avatar_style="avataaars",
                    avatar_style="🤵",

                    #avatar_style="initials",
                    #avatar_style="user",
                    #avatar_style="adventurer",
                    is_user=True,                                    
                    key=str(i) + 'data_by_user')  
                perg = st.session_state['ai_question'][i] 
                
                #Pergunta para Gemini
                language = 'pt'     # Language in which you want to convert
                # Passing the text and language to the engine,here we have marked slow=False. Which tells the module that the converted audio should have a high speed
                #myobj1 = gTTS(text="PERGUNTA: " + str(perg), lang=language, slow=False)
                myobj1 = gTTS(text= perg, lang=language, slow=False)
                name1 = "perg" + str(i) + ".mp3"
                myobj1.save(name1)   #Saving the converted audio in a mp3 file
                # Playing the converted file
                audio_file = open(name1, 'rb')
                audio_bytes = audio_file.read()
                st.audio(audio_bytes, format='audio/ogg',start_time=0)
        
                # This function displays user ai_answer
                message(st.session_state["ai_answer"][i],
                    key=str(i),
                    #avatar_style="icons"
                    avatar_style="bottts"
                    )            
                resp = st.session_state["ai_answer"][i]
                resp = resp.replace("*", "")
                resp = resp.replace("**", "")
                
                #RESPOSTA Gemini
                # Passing the text and language to the engine,here we have marked slow=False. Which tells the module that the converted audio should have a high speed
                myobj2 = gTTS(text=resp, lang=language, slow=False)
                name2 = "resp" + str(i) + ".mp3"
                myobj2.save(name2)   #Saving the converted audio in a mp3 file
                # Playing the converted file
                audio_file2 = open(name2, 'rb')
                audio_bytes2 = audio_file2.read()
                st.audio(audio_bytes2, format='audio/ogg',start_time=0) 
with tab3: 
    if GeraGrafico:
        plt.imshow(NuvemPalavras);
        plt.imshow(NuvemPalavras, interpolation='bilinear')
        plt.axis("off")
        imagem = plt.show()
        NuvemPalavras.to_file("Nuvem_Palavras.png")
        st.pyplot(imagem) #Este método faz exibirt a nuvem de palavras
        st.set_option('deprecation.showPyplotGlobalUse', False)
        Nuvem = str(NuvemPalavras)
        wordlist = Nuvem.split()
        st.write(Nuvem)
        wordfreq = []
        for w in wordlist:
            wordfreq.append(wordlist.count(w))
        chart_data = pd.DataFrame(wordfreq,wordlist)
        st.bar_chart(chart_data)


#st.image(ArqPNG, width=600, caption='Label da Figura')
st.sidebar.divider()
Rodape = '<p style="font-weight: bolder; color:white; font-size: 16px;">Desenvolvido por Massaki de O. Igarashi (Integrante  do Grupo de Pesquisa em Direito, Inovação e Tecnologia - GPDIT) | Plataforma suportada pela API Google Gemini</p>'
st.sidebar.markdown(Rodape, unsafe_allow_html=True)
st.sidebar.markdown(mystyle1, unsafe_allow_html=True) 

